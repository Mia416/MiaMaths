from docx import *
from docx.shared import Pt
import random
import operator
from docx.shared import Inches





class M_Doc:

    def _init_(self,document):
        print("init")

    def Create_Doc(self,document):
        #document = Document()
        print("create")

    def Add_Process(self,document,content):
        Tstyle = document.styles['Normal']
        Tstyle.font.size = Pt(18)
        Tstyle.font.name = 'Calibri'
        p = document.add_paragraph()
        p.style = Tstyle
        p.add_run(content)
        #print(content)
    def Add_Picture(self,document,picpath):
        document.add_picture(picpath, width=Inches(10.25))


    def Save_Doc(self,document,name):
        document.save(name)



class Gen_Data:
    def _init_(self):
        print("init")

    def Generate_Int(self,n,m):
        return random.randint(int(n),int(m))
    def Generate_IntRange(self,n,m,step):
        return random.randrange(int(n),int(m),int(step))
    def Generate_Decimal(self,n,m,t):
        return (round(random.uniform(float(n),float(m)),int(t)))
